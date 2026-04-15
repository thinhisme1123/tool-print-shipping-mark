import streamlit as st
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font as XLFont, Border, Side
from openpyxl.worksheet.pagebreak import Break
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
import io

st.set_page_config(page_title="Shipping Mark Tool V3.4 - FINAL", layout="centered")

st.title("📦 Trình Tạo Shipping Mark Cực Đại (6 Tem/Trang)")
st.write("Cập nhật V3.4: Tối ưu 6 tem/trang A4. Chữ siêu to, lấp đầy trang giấy, có khoảng cách cắt.")

uploaded_file = st.file_uploader("Tải file PKL lên", type=["xlsx", "csv"])

if uploaded_file is not None:
    try:
        if uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file, sheet_name='PKL', header=13)
        else:
            df = pd.read_csv(uploaded_file, header=13)

        df = df.dropna(subset=['P/NO'], how='all')
        packages_data = []
        current_desc, current_unit = "", ""
        
        try: total_packages = int(pd.to_numeric(df['Material code'], errors='coerce').max())
        except: total_packages = "..."

        for idx, row in df.iterrows():
            try: p_start, p_end = int(row['P/NO']), int(row['Material code'])
            except: continue

            desc = str(row['DESCRIPTION '])
            if desc.strip() and desc != 'nan': current_desc = desc
            unit = str(row['UNIT'])
            if unit.strip() and unit != 'nan': current_unit = unit
                
            qty = str(row['Unnamed: 10']).replace('.0', '')
            nw = f"{float(row['Net Weight\n(KGM)']):.2f}".rstrip('0').rstrip('.')
            gw = f"{float(row['Gross Weight\n(KGM)']):.2f}".rstrip('0').rstrip('.')

            for p in range(p_start, p_end + 1):
                data = {
                    "package_no": f"{p}/{total_packages}",
                    "item": current_desc, "qty": f"{qty} {current_unit}",
                    "nw": f"{nw} Kg", "gw": f"{gw} Kg"
                }
                packages_data.append(data)

        st.success(f"Đã xử lý {len(packages_data)} kiện hàng.")

        # --- PHẦN 1: XUẤT EXCEL 6 TEM SIÊU TO ---
        excel_out = io.BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50
        
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        r_idx, c_idx = 1, 1
        row_count = 0
        
        for d in packages_data:
            cell = ws.cell(row=r_idx, column=c_idx)
            cell.value = f"P.No: {d['package_no']}\nItem: {d['item']}\nQ.ty: {d['qty']}\nN.W: {d['nw']}\nG.W: {d['gw']}"
            # TĂNG PHÔNG CHỮ LÊN 22 CHO RÕ NÉT
            cell.font = XLFont(name='Arial', size=22, bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical='center')
            cell.border = thin_border
            
            # TĂNG CHIỀU CAO LÊN 300 ĐỂ LẤP ĐẦY TRANG (3 HÀNG)
            ws.row_dimensions[r_idx].height = 250 
            
            c_idx += 1
            if c_idx > 2: 
                c_idx = 1
                row_count += 1
                
                if row_count == 3: # Đủ 3 hàng (6 tem) -> Ngắt trang
                    ws.row_breaks.append(Break(id=r_idx))
                    row_count = 0
                    r_idx += 1
                else:
                    r_idx += 1
                    ws.row_dimensions[r_idx].height = 35 # Khoảng cách đường cắt
                    r_idx += 1
                    
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_margins.left, ws.page_margins.right = 0.2, 0.2
        ws.page_margins.top, ws.page_margins.bottom = 0.2, 0.2
        wb.save(excel_out)

        # --- PHẦN 2: XUẤT WORD 6 TEM SIÊU TO ---
        word_out = io.BytesIO()
        doc = Document()
        section = doc.sections[0]
        section.top_margin, section.bottom_margin = Cm(0.5), Cm(0.5)
        section.left_margin, section.right_margin = Cm(1), Cm(1)

        def chunker(seq, size):
            return (seq[pos:pos + size] for pos in range(0, len(seq), size))

        for page_idx, page_data in enumerate(chunker(packages_data, 6)):
            if page_idx > 0: doc.add_page_break()

            num_data_rows = len(page_data) // 2 + (len(page_data) % 2 > 0)
            total_rows = num_data_rows * 2 - 1 if num_data_rows > 0 else 0
            table = doc.add_table(rows=total_rows, cols=2)
            table.style = 'Table Grid'

            for r_i, r_obj in enumerate(table.rows):
                if r_i % 2 == 1: r_obj.height = Cm(0.7) # Khoảng trống cắt
                else: r_obj.height = Cm(8.5) # Chiều cao mỗi tem cực đại

            for i, d in enumerate(page_data):
                data_row = (i // 2) * 2 
                col = i % 2
                cell = table.cell(data_row, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.add_run(f"P.No: {d['package_no']}").bold = True
                p.add_run(f"\nItem: {d['item']}\nQ.ty: {d['qty']}\nN.W: {d['nw']}\nG.W: {d['gw']}")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name, run.font.size = 'Arial', Pt(18)

        doc.save(word_out)

        st.info("Bản V3.4: Chữ siêu to, lấp đầy 3 hàng/trang A4.")
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("📥 Tải EXCEL (6 Tem Cực Đại)", excel_out.getvalue(), "Final_Max_Marks.xlsx")
        with col2:
            st.download_button("📥 Tải WORD (6 Tem Cực Đại)", word_out.getvalue(), "Final_Max_Marks.docx")

    except Exception as e:
        st.error(f"Lỗi: {e}")